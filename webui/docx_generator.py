"""
DOCX Generator for professional document creation
Creates formatted Word documents from enriched transcript sections
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional


class DocxGenerator:
    """Generate professional DOCX documents from transcript data"""

    def __init__(self, language: str = 'en'):
        self.doc = Document()
        self.language = language.lower()
        self._setup_styles()
        self._set_document_language()

    def _setup_styles(self):
        """Configure document styles - optimized for readability"""
        styles = self.doc.styles

        # Heading 1 style - Georgia for elegant serif headings
        if 'CustomHeading1' not in styles:
            heading1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.base_style = styles['Heading 1']
            heading1.font.name = 'Georgia'
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 0, 0)  # Black for better contrast
            # Spacing after heading
            heading1.paragraph_format.space_after = Pt(12)
            heading1.paragraph_format.space_before = Pt(18)

        # Heading 2 style - Georgia for section titles
        if 'CustomHeading2' not in styles:
            heading2 = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.base_style = styles['Heading 2']
            heading2.font.name = 'Georgia'
            heading2.font.size = Pt(13)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(0, 0, 0)
            # Good spacing for section separation
            heading2.paragraph_format.space_after = Pt(8)
            heading2.paragraph_format.space_before = Pt(16)

        # Body text style - Calibri for excellent readability
        if 'CustomBody' not in styles:
            body = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body.base_style = styles['Normal']
            body.font.name = 'Calibri'
            body.font.size = Pt(11)
            # Line spacing for comfortable reading
            body.paragraph_format.line_spacing = 1.15

    def _set_document_language(self):
        """Set document language for spell checking"""
        # Map common language codes to Word language identifiers
        lang_map = {
            'fr': 'fr-FR',
            'en': 'en-US',
            'es': 'es-ES',
            'de': 'de-DE',
            'it': 'it-IT',
            'pt': 'pt-PT',
            'nl': 'nl-NL',
            'pl': 'pl-PL',
            'ru': 'ru-RU',
            'ja': 'ja-JP',
            'zh': 'zh-CN',
            'ar': 'ar-SA',
        }

        lang_code = lang_map.get(self.language, 'en-US')

        # Set language for the document's default style
        try:
            from docx.oxml.shared import OxmlElement, qn

            # Get the document's styles element
            styles_element = self.doc.styles.element

            # Create or update the document defaults
            doc_defaults = styles_element.find(qn('w:docDefaults'))
            if doc_defaults is None:
                doc_defaults = OxmlElement('w:docDefaults')
                styles_element.insert(0, doc_defaults)

            # Set run properties defaults
            rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
            if rPrDefault is None:
                rPrDefault = OxmlElement('w:rPrDefault')
                doc_defaults.append(rPrDefault)

            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDefault.append(rPr)

            # Set language
            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                rPr.append(lang)

            lang.set(qn('w:val'), lang_code)
            lang.set(qn('w:eastAsia'), lang_code)
            lang.set(qn('w:bidi'), lang_code)

        except Exception as e:
            print(f"[DOCX] Warning: Could not set document language: {e}")

    def add_document_title(self, title: str):
        """
        Add document title at top of first page (no separate cover page)

        Args:
            title: Document title
        """
        # Title as Heading 1 at top of document
        title_heading = self.doc.add_heading(title, level=1)
        title_heading.style = 'CustomHeading1'
        title_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-aligned like a book
        # No extra paragraph - spacing is controlled by style

    def add_summary(self, summary_text: str):
        """
        Add executive summary section

        Args:
            summary_text: Summary content
        """
        # Summary heading
        heading = self.doc.add_heading('Résumé', level=1)
        heading.style = 'CustomHeading1'

        # Summary content
        para = self.doc.add_paragraph(summary_text)
        para.style = 'CustomBody'

        self.doc.add_page_break()

    def add_section(self, section: Dict):
        """
        Add enriched section to document

        Args:
            section: Dict with title and content (naturally structured by LLM)
        """
        # Section title (Heading 2)
        title = section.get('title', 'Section sans titre')
        heading = self.doc.add_heading(title, level=2)
        heading.style = 'CustomHeading2'

        # Main content - split into paragraphs
        # The LLM will naturally structure the content with its own organization
        content = section.get('content', '')
        if content:
            # Split content by double newlines (paragraph breaks) or single newlines
            paragraphs = content.split('\n\n') if '\n\n' in content else content.split('\n')

            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:  # Skip empty paragraphs
                    para = self.doc.add_paragraph(para_text)
                    para.style = 'CustomBody'
                    # Justify text and add first line indent
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.first_line_indent = Inches(0.3)
                    # Add spacing after paragraph
                    para.paragraph_format.space_after = Pt(6)

            # No extra paragraph - spacing between sections is controlled by Heading 2 space_before

    def save(self, output_path: str):
        """
        Save document to file

        Args:
            output_path: Path to save DOCX file
        """
        self.doc.save(output_path)
        print(f"[DOCX] Document saved: {output_path}")

    def _format_doc_type(self, doc_type: str) -> str:
        """Format document type for display"""
        type_map = {
            'course': 'Notes de cours',
            'meeting': 'Compte-rendu de réunion',
            'conference': 'Résumé de conférence',
            'interview': 'Transcription d\'interview',
            'other': 'Document'
        }
        return type_map.get(doc_type, 'Document')


def create_document(title: str, doc_type: str, sections: List[Dict],
                   summary: Optional[str] = None, metadata: Optional[Dict] = None) -> Document:
    """
    Create complete DOCX document from transcript data

    Args:
        title: Document title
        doc_type: Type of document (course, meeting, etc.)
        sections: List of enriched sections
        summary: Optional executive summary
        metadata: Optional metadata dict (date, duration, etc.)

    Returns:
        Document object
    """
    generator = DocxGenerator()

    # Simple title at top (no separate cover page)
    generator.add_document_title(title)

    # No table of contents or summary - keep it simple and clean
    # Just add the enriched content sections directly

    # Sections
    for section in sections:
        generator.add_section(section)

    return generator.doc


def generate_docx_file(title: str, doc_type: str, sections: List[Dict],
                       summary: Optional[str], metadata: Dict, output_path: str,
                       language: str = 'en') -> str:
    """
    Generate and save DOCX file

    Args:
        title: Document title
        doc_type: Type of document
        sections: List of enriched sections
        summary: Optional summary text
        metadata: Metadata dict
        output_path: Path to save file
        language: Document language code (fr, en, es, etc.)

    Returns:
        Path to saved file
    """
    print(f"[DOCX] Generating document: {title} (language: {language})")

    generator = DocxGenerator(language=language)

    # Simple title at top (no separate cover page)
    generator.add_document_title(title)

    # No table of contents or summary - keep it clean and simple

    for i, section in enumerate(sections, 1):
        print(f"[DOCX] Adding section {i}/{len(sections)}: {section.get('title', 'Untitled')}")
        generator.add_section(section)

    generator.save(output_path)

    return output_path
