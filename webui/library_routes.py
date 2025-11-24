"""
Library routes for document management, statistics, and job history
"""
from flask import Blueprint, render_template, request, jsonify, send_file, abort, flash, redirect, url_for
from flask_login import login_required, current_user
from sqlalchemy import func, desc, asc, and_, or_
from datetime import datetime, timedelta
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
import os
from pathlib import Path

from database import SessionLocal
from models import Document, Job, User
from auth import verify_password
from docx import Document as DocxDocument
from weasyprint import HTML
import tempfile
import zipfile
import shutil

library_bp = Blueprint('library', __name__)

# Token serializer for secure download links
serializer = None

def init_serializer(secret_key):
    """Initialize the URL serializer with app secret key"""
    global serializer
    serializer = URLSafeTimedSerializer(secret_key)


def set_docx_language(doc, language_code):
    """Set document language for spell checking"""
    if not language_code:
        return

    # Map common language codes to Word language identifiers
    lang_map = {
        'fr': 'fr-FR',
        'en': 'en-US',
        'es': 'es-ES',
        'de': 'de-DE',
        'it': 'it-IT',
        'pt': 'pt-PT',
        'nl': 'nl-NL',
        'pl': 'pl-PL',
        'ru': 'ru-RU',
        'ja': 'ja-JP',
        'zh': 'zh-CN',
        'ar': 'ar-SA'
    }

    lang_code = lang_map.get(language_code.lower(), 'en-US')

    try:
        from docx.oxml.shared import OxmlElement, qn

        # Set language in document settings
        settings = doc.settings
        if hasattr(settings, 'element'):
            settings_elem = settings.element

            # Create or find lang defaults element
            lang_defaults = settings_elem.find(qn('w:themeFontLang'))
            if lang_defaults is None:
                lang_defaults = OxmlElement('w:themeFontLang')
                settings_elem.append(lang_defaults)

            lang_defaults.set(qn('w:val'), lang_code)

        # Set language for each paragraph's runs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                rPr = run._element.get_or_add_rPr()
                lang = rPr.find(qn('w:lang'))
                if lang is None:
                    lang = OxmlElement('w:lang')
                    rPr.append(lang)

                lang.set(qn('w:val'), lang_code)

    except Exception as e:
        print(f"[EXPORT] Warning: Could not set document language: {e}")


# ============================================================================
# LIBRARY - Document Management
# ============================================================================

@library_bp.route('/library')
@login_required
def library():
    """Main library page with documents list"""
    db = SessionLocal()
    try:
        # Get filter parameters
        search = request.args.get('search', '').strip()
        doc_type = request.args.get('type', '')
        language = request.args.get('language', '')
        mode = request.args.get('mode', '')
        favorites_only = request.args.get('favorites', '') == 'true'
        sort_by = request.args.get('sort', 'date_desc')
        page = int(request.args.get('page', 1))
        per_page = 20

        # Build query
        query = db.query(Document).filter(Document.user_id == current_user.id)

        # Apply filters
        if search:
            query = query.filter(Document.title.ilike(f'%{search}%'))

        if doc_type:
            query = query.filter(Document.document_type == doc_type)

        if language:
            query = query.filter(Document.language == language)

        if mode:
            query = query.filter(Document.mode == mode)

        if favorites_only:
            query = query.filter(Document.is_favorite == True)

        # Apply sorting
        if sort_by == 'date_desc':
            query = query.order_by(desc(Document.created_at))
        elif sort_by == 'date_asc':
            query = query.order_by(asc(Document.created_at))
        elif sort_by == 'title_asc':
            query = query.order_by(asc(Document.title))
        elif sort_by == 'title_desc':
            query = query.order_by(desc(Document.title))
        elif sort_by == 'size_desc':
            query = query.order_by(desc(Document.file_size_bytes))
        elif sort_by == 'size_asc':
            query = query.order_by(asc(Document.file_size_bytes))

        # Paginate
        total_docs = query.count()
        documents = query.limit(per_page).offset((page - 1) * per_page).all()

        # Calculate storage stats
        storage_stats = calculate_storage_stats(db, current_user.id)

        # Get unique values for filters
        doc_types = db.query(Document.document_type).filter(
            Document.user_id == current_user.id,
            Document.document_type.isnot(None)
        ).distinct().all()
        doc_types = [t[0] for t in doc_types if t[0]]

        languages = db.query(Document.language).filter(
            Document.user_id == current_user.id,
            Document.language.isnot(None)
        ).distinct().all()
        languages = [l[0] for l in languages if l[0]]

        # Get all unique tags
        all_tags = set()
        for doc in db.query(Document).filter(Document.user_id == current_user.id).all():
            if doc.tags:
                all_tags.update(doc.tags)
        all_tags = sorted(list(all_tags))

        return render_template(
            'library.html',
            documents=documents,
            total_docs=total_docs,
            page=page,
            per_page=per_page,
            total_pages=(total_docs + per_page - 1) // per_page,
            storage_stats=storage_stats,
            doc_types=doc_types,
            languages=languages,
            all_tags=all_tags,
            # Current filters
            current_search=search,
            current_type=doc_type,
            current_language=language,
            current_mode=mode,
            favorites_only=favorites_only,
            sort_by=sort_by
        )
    finally:
        db.close()


@library_bp.route('/api/documents/<int:doc_id>/toggle-favorite', methods=['POST'])
@login_required
def toggle_favorite(doc_id):
    """Toggle favorite status of a document"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == current_user.id
        ).first()

        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        doc.is_favorite = not doc.is_favorite
        db.commit()

        return jsonify({'success': True, 'is_favorite': doc.is_favorite})
    finally:
        db.close()


@library_bp.route('/api/documents/<int:doc_id>/tags', methods=['POST'])
@login_required
def update_tags(doc_id):
    """Update tags for a document"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == current_user.id
        ).first()

        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        tags = request.json.get('tags', [])
        mode = request.json.get('mode', 'replace')  # 'replace' or 'add'

        # Clean and validate new tags
        new_tags = [tag.strip() for tag in tags if tag.strip()]

        if mode == 'add':
            # Add new tags to existing ones (avoid duplicates)
            current_tags = doc.tags if doc.tags else []
            combined_tags = list(set(current_tags + new_tags))
            doc.tags = combined_tags[:10]  # Max 10 tags
        else:
            # Replace all tags
            doc.tags = new_tags[:10]

        db.commit()

        return jsonify({'success': True, 'tags': doc.tags})
    finally:
        db.close()


@library_bp.route('/api/documents/<int:doc_id>/delete', methods=['POST'])
@login_required
def delete_document(doc_id):
    """Delete a document"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == current_user.id
        ).first()

        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        # Delete file
        if os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except Exception as e:
                print(f"[LIBRARY] Error deleting file {doc.file_path}: {e}")

        # Delete from DB
        db.delete(doc)
        db.commit()

        return jsonify({'success': True})
    finally:
        db.close()


@library_bp.route('/api/library/delete-all', methods=['POST'])
@login_required
def delete_library():
    """Delete all documents for current user (requires password confirmation)"""
    password = request.json.get('password', '')

    if not verify_password(current_user.password_hash, password):
        return jsonify({'error': 'Mot de passe incorrect'}), 401

    db = SessionLocal()
    try:
        # Get all user documents
        documents = db.query(Document).filter(Document.user_id == current_user.id).all()

        deleted_count = 0
        for doc in documents:
            # Delete file
            if os.path.exists(doc.file_path):
                try:
                    os.remove(doc.file_path)
                    deleted_count += 1
                except Exception as e:
                    print(f"[LIBRARY] Error deleting file {doc.file_path}: {e}")

            # Delete from DB
            db.delete(doc)

        db.commit()

        return jsonify({'success': True, 'deleted_count': deleted_count})
    finally:
        db.close()


@library_bp.route('/library/download/<token>')
def download_document(token):
    """Download a document using a signed token"""
    try:
        # Verify token (24h expiry)
        data = serializer.loads(token, max_age=86400)  # 24 hours
        doc_id = data['doc_id']
        user_id = data['user_id']

        # Verify user (even if not logged in, token contains user_id)
        # If user is logged in, verify it matches
        if current_user.is_authenticated and current_user.id != user_id:
            abort(403)

        db = SessionLocal()
        try:
            doc = db.query(Document).filter(
                Document.id == doc_id,
                Document.user_id == user_id
            ).first()

            if not doc or not os.path.exists(doc.file_path):
                abort(404)

            # Get filename
            filename = f"{doc.title}.{doc.file_path.split('.')[-1]}"

            return send_file(
                doc.file_path,
                as_attachment=True,
                download_name=filename
            )
        finally:
            db.close()

    except SignatureExpired:
        flash('Le lien de téléchargement a expiré', 'error')
        return redirect(url_for('library.library'))
    except BadSignature:
        flash('Lien de téléchargement invalide', 'error')
        return redirect(url_for('library.library'))


@library_bp.route('/api/documents/<int:doc_id>/download-link', methods=['POST'])
@login_required
def generate_download_link(doc_id):
    """Generate a temporary download link for a document"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == current_user.id
        ).first()

        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        # Generate signed token
        token = serializer.dumps({'doc_id': doc.id, 'user_id': current_user.id})
        download_url = url_for('library.download_document', token=token, _external=True)

        return jsonify({'success': True, 'download_url': download_url})
    finally:
        db.close()


@library_bp.route('/library/export/<int:doc_id>/<format>')
@login_required
def export_document(doc_id, format):
    """Export document in specified format (docx, pdf, markdown)"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == current_user.id
        ).first()

        if not doc or not os.path.exists(doc.file_path):
            abort(404)

        # DOCX - direct download or create from text
        if format == 'docx':
            file_ext = os.path.splitext(doc.file_path)[1].lower()

            if file_ext == '.docx':
                # Direct download for existing DOCX
                filename = f"{doc.title}.docx"
                return send_file(
                    doc.file_path,
                    as_attachment=True,
                    download_name=filename
                )
            else:
                # Create DOCX from plain text file
                try:
                    new_doc = DocxDocument()
                    new_doc.add_heading(doc.title, 0)

                    # Read text file
                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        for line in content.split('\n'):
                            new_doc.add_paragraph(line)

                    # Set document language if available
                    if doc.language:
                        set_docx_language(new_doc, doc.language)

                    # Save to temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        new_doc.save(tmp.name)
                        tmp_path = tmp.name

                    try:
                        filename = f"{doc.title}.docx"
                        return send_file(
                            tmp_path,
                            as_attachment=True,
                            download_name=filename,
                            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                    finally:
                        if os.path.exists(tmp_path):
                            os.remove(tmp_path)

                except Exception as e:
                    print(f"[EXPORT] DOCX generation error: {e}")
                    flash('Erreur lors de la génération du DOCX', 'error')
                    return redirect(url_for('library.library'))

        # PDF - convert from DOCX or text
        elif format == 'pdf':
            try:
                # Detect file type
                file_ext = os.path.splitext(doc.file_path)[1].lower()

                # Build HTML
                html_content = '<html><head><meta charset="utf-8"><style>'
                html_content += 'body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }'
                html_content += 'h1 { color: #333; border-bottom: 2px solid #667eea; padding-bottom: 10px; }'
                html_content += 'p { margin: 10px 0; color: #333; white-space: pre-wrap; }'
                html_content += '</style></head><body>'
                html_content += f'<h1>{doc.title}</h1>'

                # Extract content based on file type
                if file_ext == '.docx':
                    # Read DOCX content
                    docx_doc = DocxDocument(doc.file_path)
                    for para in docx_doc.paragraphs:
                        if para.text.strip():
                            html_content += f'<p>{para.text}</p>'
                else:
                    # Read plain text file (txt, srt, etc.)
                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        for line in content.split('\n'):
                            if line.strip():
                                html_content += f'<p>{line}</p>'

                html_content += '</body></html>'

                # Generate PDF
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    HTML(string=html_content).write_pdf(tmp.name)
                    tmp_path = tmp.name

                try:
                    filename = f"{doc.title}.pdf"
                    return send_file(
                        tmp_path,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/pdf'
                    )
                finally:
                    # Cleanup temp file after sending
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)

            except Exception as e:
                print(f"[EXPORT] PDF generation error: {e}")
                flash('Erreur lors de la génération du PDF', 'error')
                return redirect(url_for('library.library'))

        # MARKDOWN - simple text export
        elif format == 'markdown':
            try:
                # Detect file type
                file_ext = os.path.splitext(doc.file_path)[1].lower()

                # Build Markdown
                md_content = f"# {doc.title}\n\n"

                # Extract content based on file type
                if file_ext == '.docx':
                    # Read DOCX content
                    docx_doc = DocxDocument(doc.file_path)
                    for para in docx_doc.paragraphs:
                        if para.text.strip():
                            md_content += f"{para.text}\n\n"
                else:
                    # Read plain text file (txt, srt, etc.)
                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                        md_content += f.read()

                # Create temp file
                with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.md', encoding='utf-8') as tmp:
                    tmp.write(md_content)
                    tmp_path = tmp.name

                try:
                    filename = f"{doc.title}.md"
                    return send_file(
                        tmp_path,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='text/markdown'
                    )
                finally:
                    # Cleanup temp file after sending
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)

            except Exception as e:
                print(f"[EXPORT] Markdown generation error: {e}")
                flash('Erreur lors de la génération du Markdown', 'error')
                return redirect(url_for('library.library'))

        else:
            abort(400)

    finally:
        db.close()


@library_bp.route('/library/bulk-export', methods=['POST'])
@login_required
def bulk_export():
    """Export multiple documents as a ZIP file in the specified format"""
    try:
        data = request.get_json()
        doc_ids = data.get('doc_ids', [])
        format = data.get('format', 'docx')

        if not doc_ids:
            return jsonify({'error': 'No documents selected'}), 400

        db = SessionLocal()
        try:
            # Get documents
            docs = db.query(Document).filter(
                Document.id.in_(doc_ids),
                Document.user_id == current_user.id
            ).all()

            if not docs:
                return jsonify({'error': 'No documents found'}), 404

            # Create temporary directory for ZIP
            temp_dir = tempfile.mkdtemp()
            zip_path = os.path.join(temp_dir, f'documents_{format}.zip')

            try:
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for doc in docs:
                        if not os.path.exists(doc.file_path):
                            continue

                        # Detect file type
                        file_ext = os.path.splitext(doc.file_path)[1].lower()

                        # Generate file based on format
                        if format == 'docx':
                            # For DOCX: copy directly if already DOCX, create from text otherwise
                            if file_ext == '.docx':
                                filename = f"{doc.title}.docx"
                                zipf.write(doc.file_path, filename)
                            else:
                                # Create DOCX from plain text file
                                try:
                                    new_doc = DocxDocument()
                                    new_doc.add_heading(doc.title, 0)

                                    # Read text file
                                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                                        content = f.read()
                                        # Add content as paragraphs
                                        for line in content.split('\n'):
                                            new_doc.add_paragraph(line)

                                    # Set document language if available
                                    if doc.language:
                                        set_docx_language(new_doc, doc.language)

                                    # Save to temp file
                                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                        new_doc.save(tmp.name)
                                        tmp_docx = tmp.name

                                    # Add to ZIP
                                    filename = f"{doc.title}.docx"
                                    zipf.write(tmp_docx, filename)
                                    os.remove(tmp_docx)

                                except Exception as e:
                                    print(f"[EXPORT] Error creating DOCX for {doc.title}: {e}")
                                    continue

                        elif format == 'pdf':
                            # Convert to PDF
                            try:
                                # Build HTML
                                html_content = '<html><head><meta charset="utf-8"><style>'
                                html_content += 'body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }'
                                html_content += 'h1 { color: #333; border-bottom: 2px solid #667eea; padding-bottom: 10px; }'
                                html_content += 'p { margin: 10px 0; color: #333; white-space: pre-wrap; }'
                                html_content += '</style></head><body>'
                                html_content += f'<h1>{doc.title}</h1>'

                                # Extract content based on file type
                                if file_ext == '.docx':
                                    # Read DOCX
                                    docx_doc = DocxDocument(doc.file_path)
                                    for para in docx_doc.paragraphs:
                                        if para.text.strip():
                                            html_content += f'<p>{para.text}</p>'
                                else:
                                    # Read plain text file (txt, srt, etc.)
                                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                                        content = f.read()
                                        # Split by lines and wrap in paragraphs
                                        for line in content.split('\n'):
                                            if line.strip():
                                                html_content += f'<p>{line}</p>'

                                html_content += '</body></html>'

                                # Generate PDF to temp file
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                                    HTML(string=html_content).write_pdf(tmp.name)
                                    tmp_pdf = tmp.name

                                # Add to ZIP
                                filename = f"{doc.title}.pdf"
                                zipf.write(tmp_pdf, filename)
                                os.remove(tmp_pdf)

                            except Exception as e:
                                print(f"[EXPORT] Error converting {doc.title} to PDF: {e}")
                                continue

                        elif format == 'markdown':
                            # Convert to Markdown
                            try:
                                md_content = f"# {doc.title}\n\n"

                                # Extract content based on file type
                                if file_ext == '.docx':
                                    # Read DOCX
                                    docx_doc = DocxDocument(doc.file_path)
                                    for para in docx_doc.paragraphs:
                                        if para.text.strip():
                                            md_content += f"{para.text}\n\n"
                                else:
                                    # Read plain text file (txt, srt, etc.)
                                    with open(doc.file_path, 'r', encoding='utf-8') as f:
                                        md_content += f.read()

                                # Add to ZIP
                                filename = f"{doc.title}.md"
                                zipf.writestr(filename, md_content.encode('utf-8'))

                            except Exception as e:
                                print(f"[EXPORT] Error converting {doc.title} to Markdown: {e}")
                                continue

                # Send ZIP file
                response = send_file(
                    zip_path,
                    as_attachment=True,
                    download_name=f'documents_{format}.zip',
                    mimetype='application/zip'
                )

                # Clean up temp directory after sending
                @response.call_on_close
                def cleanup():
                    try:
                        shutil.rmtree(temp_dir)
                    except:
                        pass

                return response

            except Exception as e:
                # Clean up on error
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                raise e

        finally:
            db.close()

    except Exception as e:
        print(f"[EXPORT] Bulk export error: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================================================
# JOBS HISTORY
# ============================================================================

@library_bp.route('/jobs')
@login_required
def jobs_history():
    """Job history page with filters"""
    db = SessionLocal()
    try:
        # Get filter parameters
        status_filter = request.args.get('status', '')
        mode_filter = request.args.get('mode', '')
        page = int(request.args.get('page', 1))
        per_page = 50

        # Build query
        query = db.query(Job).filter(Job.user_id == current_user.id)

        # Apply filters
        if status_filter:
            query = query.filter(Job.status == status_filter)

        if mode_filter:
            query = query.filter(Job.mode == mode_filter)

        # Sort by date desc
        query = query.order_by(desc(Job.created_at))

        # Paginate
        total_jobs = query.count()
        jobs = query.limit(per_page).offset((page - 1) * per_page).all()

        # Calculate stats
        total_completed = db.query(func.count(Job.id)).filter(
            Job.user_id == current_user.id,
            Job.status == 'completed'
        ).scalar()

        total_errors = db.query(func.count(Job.id)).filter(
            Job.user_id == current_user.id,
            Job.status == 'error'
        ).scalar()

        return render_template(
            'jobs_history.html',
            jobs=jobs,
            total_jobs=total_jobs,
            total_completed=total_completed,
            total_errors=total_errors,
            page=page,
            per_page=per_page,
            total_pages=(total_jobs + per_page - 1) // per_page,
            status_filter=status_filter,
            mode_filter=mode_filter
        )
    finally:
        db.close()


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def calculate_storage_stats(db, user_id):
    """Calculate storage statistics for a user"""
    total_size = db.query(func.sum(Document.file_size_bytes)).filter(
        Document.user_id == user_id
    ).scalar() or 0

    total_docs = db.query(func.count(Document.id)).filter(
        Document.user_id == user_id
    ).scalar()

    # Get user storage limit
    user = db.query(User).get(user_id)
    storage_limit = user.storage_limit_bytes if user else 2 * 1024 * 1024 * 1024  # 2GB default

    percentage = (total_size / storage_limit * 100) if storage_limit > 0 else 0

    return {
        'total_size': total_size,
        'total_size_mb': total_size / (1024 * 1024),
        'total_size_gb': total_size / (1024 * 1024 * 1024),
        'storage_limit': storage_limit,
        'storage_limit_gb': storage_limit / (1024 * 1024 * 1024),
        'percentage': min(percentage, 100),
        'total_docs': total_docs
    }


